"""
סקריפט להמרת קובץ Word עם נתוני רישוי עסקים לפורמט CSV
משתמש ב-Polars לעיבוד יעיל של הנתונים
"""

import os
import re
from docx import Document
import polars as pl
from typing import List, Dict, Any


def read_docx_file(file_path: str) -> str:
    """
    קריאת קובץ Word והחזרת התוכן כטקסט
    """
    try:
        doc = Document(file_path)
        full_text = []
        
        # קריאת כל הפסקאות
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text.strip())
        
        # קריאת טבלאות אם יש
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        return "\n".join(full_text)
    
    except Exception as e:
        print(f"שגיאה בקריאת הקובץ: {e}")
        return ""


def extract_licensing_requirements(text: str) -> List[Dict[str, Any]]:
    """
    חילוץ דרישות רישוי מהטקסט
    מחזיר רשימה של דיקטים עם הדרישות
    """
    requirements = []
    lines = text.split('\n')
    
    current_requirement = {}
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # זיהוי כותרות או דרישות חדשות
        if any(keyword in line for keyword in ['רישוי', 'היתר', 'דרישה', 'חובה', 'צריך']):
            # שמירת הדרישה הקודמת אם קיימת
            if current_requirement:
                requirements.append(current_requirement.copy())
                current_requirement = {}
            
            current_requirement['title'] = line
            current_requirement['description'] = line
            current_requirement['authority'] = ''
            current_requirement['area_requirements'] = ''
            current_requirement['capacity_requirements'] = ''
            current_requirement['special_requirements'] = ''
            current_requirement['priority'] = 'medium'
            current_requirement['cost_estimate'] = ''
            current_requirement['processing_time'] = ''
        
        # חיפוש מספרים שיכולים להצביע על שטח או תפוסה
        area_match = re.search(r'(\d+)\s*(מ"ר|מטר|שטח)', line)
        if area_match and current_requirement:
            current_requirement['area_requirements'] = area_match.group(0)
        
        capacity_match = re.search(r'(\d+)\s*(מקומות|אנשים|לקוחות)', line)
        if capacity_match and current_requirement:
            current_requirement['capacity_requirements'] = capacity_match.group(0)
        
        # זיהוי רשויות
        if any(authority in line for authority in ['עירייה', 'משטרה', 'כבאות', 'משרד', 'רשות']):
            if current_requirement:
                current_requirement['authority'] = line
        
        # זיהוי מאפיינים מיוחדים
        if any(feature in line for feature in ['גז', 'בשר', 'משלוחים', 'אלכוהול', 'חוץ']):
            if current_requirement:
                current_requirement['special_requirements'] += line + '; '
        
        # זיהוי עלויות
        cost_match = re.search(r'(\d+)\s*(שקל|ש"ח|nis)', line)
        if cost_match and current_requirement:
            current_requirement['cost_estimate'] = cost_match.group(0)
        
        # זיהוי זמני טיפול
        time_match = re.search(r'(\d+)\s*(יום|שבוע|חודש|שנה)', line)
        if time_match and current_requirement:
            current_requirement['processing_time'] = time_match.group(0)
    
    # הוספת הדרישה האחרונה
    if current_requirement:
        requirements.append(current_requirement)
    
    return requirements


def create_polars_dataframe(requirements: List[Dict[str, Any]]) -> pl.DataFrame:
    """
    יצירת DataFrame של Polars מהדרישות
    """
    if not requirements:
        print("לא נמצאו דרישות לעיבוד")
        return pl.DataFrame()
    
    # יצירת DataFrame
    df = pl.DataFrame(requirements)
    
    # ניקוי ועיבוד נתונים
    df = df.with_columns([
        # ניקוי טקסטים
        pl.col("title").str.strip_chars(),
        pl.col("description").str.strip_chars(),
        pl.col("authority").str.strip_chars(),
        pl.col("special_requirements").str.strip_chars().str.replace_all("; $", ""),
        
        # הוספת ID ייחודי
        pl.int_range(pl.len()).alias("requirement_id")
    ])
    
    return df


def save_to_csv(df: pl.DataFrame, output_path: str):
    """
    שמירת DataFrame ל-CSV
    """
    try:
        df.write_csv(output_path, separator=',')
        print(f"הקובץ נשמר בהצלחה: {output_path}")
        
        # הדפסת סטטיסטיקות
        print(f"נמצאו {len(df)} דרישות רישוי")
        print(f"עמודות: {', '.join(df.columns)}")
        
    except Exception as e:
        print(f"שגיאה בשמירת הקובץ: {e}")


def main():
    """
    הפונקציה הראשית - המרת הקובץ מ-DOCX ל-CSV
    """
    # נתיבי קבצים
    docx_file = r"C:\Users\ah147\Desktop\A-impact\18-07-2022_4.2A.docx"
    csv_output = r"C:\Users\ah147\Desktop\A-impact\a-impact\data_processing\licensing_requirements.csv"
    
    print("מתחיל המרת קובץ דרישות הרישוי...")
    
    # בדיקת קיום הקובץ
    if not os.path.exists(docx_file):
        print(f"שגיאה: הקובץ לא נמצא: {docx_file}")
        return
    
    # קריאת הקובץ
    print("קורא את קובץ ה-Word...")
    text_content = read_docx_file(docx_file)
    
    if not text_content:
        print("שגיאה: לא ניתן לקרוא את תוכן הקובץ")
        return
    
    print(f"נקראו {len(text_content)} תווים מהקובץ")
    
    # חילוץ דרישות
    print("מחלץ דרישות רישוי...")
    requirements = extract_licensing_requirements(text_content)
    
    if not requirements:
        print("לא נמצאו דרישות בקובץ")
        return
    
    # יצירת DataFrame
    print("יוצר DataFrame עם Polars...")
    df = create_polars_dataframe(requirements)
    
    # שמירה ל-CSV
    print("שומר ל-CSV...")
    save_to_csv(df, csv_output)
    
    # הצגת דוגמה מהנתונים
    print("\nדוגמה מהנתונים:")
    print(df.head())


if __name__ == "__main__":
    main()
